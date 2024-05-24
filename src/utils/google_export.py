from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
import io
from PyPDF2 import PdfReader
import docx2txt
import os
import docx
import csv
import openpyxl

SCOPES = ['https://www.googleapis.com/auth/drive.readonly', 'https://www.googleapis.com/auth/drive.metadata.readonly']
     
def extract_text_from_drawing(drawing):
    text = ""
    for shape in drawing.shapes:
        if shape.has_text_frame:
            for paragraph in shape.text_frame.paragraphs:
                text += paragraph.text + "\n"
    return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    
    # Extract text from drawing objects
    for rel in doc.part.rels.values():
        if "drawing" in rel.reltype:
            drawing = rel.target_part.blob
            drawing_text = extract_text_from_drawing(drawing)
            text += drawing_text
    
    return text

def extract_text_from_excel(workbook):
    text = ''
    for sheet in workbook:
        for row in sheet.iter_rows():
            row_text = ' '.join(str(cell.value) for cell in row)
            text += row_text + '\n'
    return text

def extract_text_from_csv(csv_content):
    text = ''
    csv_reader = csv.reader(io.StringIO(csv_content))
    for row in csv_reader:
        row_text = ' '.join(str(cell) for cell in row)
        text += row_text + '\n'
    return text

def export_summer_camp_text(creds):
    drive_service = build('drive', 'v3', credentials=creds)
    
    # Search for files containing 'summer camp' in the name
    query = "fullText contains 'summer' or name contains 'summer'"
    
    try:
        results = drive_service.files().list(q=query, fields="nextPageToken, files(id, name, mimeType, parents)").execute()
        items = results.get('files', [])

        if not items:
            print('No files found containing summer camp information.')
            return
        
        # Create the data/google directory if it doesn't exist
        os.makedirs('data/google', exist_ok=True)
        
        for item in items:
            file_id = item['id']
            file_name = item['name']
            mime_type = item['mimeType']
            
            # Get the file path
            file_path = get_file_path(drive_service, file_id)
            
            if mime_type == 'application/vnd.google-apps.document':
                # Download Google Doc as Word document
                export_mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                request = drive_service.files().export_media(fileId=file_id, mimeType=export_mime_type)
                file_content = io.BytesIO(request.execute())
                file_content = extract_text_from_docx(file_content)  # Process the Word document

            elif mime_type == 'application/pdf':
                # Download PDF file and extract text
                request = drive_service.files().get_media(fileId=file_id)
                file_content = io.BytesIO(request.execute())
                pdf_reader = PdfReader(file_content)
                file_content = "\n".join(page.extract_text() for page in pdf_reader.pages)

            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Download Word document and extract text
                request = drive_service.files().get_media(fileId=file_id)
                file_content = io.BytesIO(request.execute())
                file_content = extract_text_from_docx(file_content)  # Process the Word document

            elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                # Download Excel file and extract text
                request = drive_service.files().get_media(fileId=file_id)
                file_content = io.BytesIO(request.execute())
                workbook = openpyxl.load_workbook(file_content)
                file_content = extract_text_from_excel(workbook)

            elif mime_type == 'text/csv':
                # Download CSV file and extract text
                request = drive_service.files().get_media(fileId=file_id)
                file_content = request.execute().decode('utf-8')
                file_content = extract_text_from_csv(file_content)

            elif mime_type == 'application/vnd.google-apps.spreadsheet':
                # Download Google Sheets as CSV and extract text
                export_mime_type = 'text/csv'
                request = drive_service.files().export_media(fileId=file_id, mimeType=export_mime_type)
                file_content = request.execute().decode('utf-8')
                file_content = extract_text_from_csv(file_content)

            else:
                print(f'Skipping unsupported file type: {file_name}')
                continue

            print(f'Exporting text from: {file_name}')
            print(f'Google Drive path: {file_path}')
            file_name = file_name.replace('/', '_')  # Replace forward slashes with underscores
            file_path = os.path.join('data', 'google', f'{file_name}.txt')
            with open(file_path, 'w') as f:
                f.write(file_content)
            
    except HttpError as error:
        print(f'An error occurred: {error}')

def get_file_path(drive_service, file_id):
    file_path = []
    while True:
        file = drive_service.files().get(fileId=file_id, fields='id, name, parents').execute()
        file_path.append(file['name'])
        if 'parents' in file:
            file_id = file['parents'][0]
        else:
            break
    file_path.reverse()
    return '/'.join(file_path)

# Set the path to the credentials.json file
credentials_path = 'credentials.json'

# Authenticate and get credentials (not shown)
creds = None
if creds and creds.expired and creds.refresh_token:
    creds.refresh(Request())
else:
    flow = InstalledAppFlow.from_client_secrets_file(credentials_path, SCOPES)
    creds = flow.run_local_server(port=0)
    with open('token.json', 'w') as token:
        token.write(creds.to_json())

export_summer_camp_text(creds)